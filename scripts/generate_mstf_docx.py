"""Generate MSTF architecture DOCX onto the user Desktop.

Includes accurate Existing vs Proposed architecture diagrams (Mermaid-equivalent).
"""
from __future__ import annotations

import os

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_header_row(row):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=(255, 255, 255))


def _box(ax, x, y, w, h, lines, fc, ec, fontsize=11):
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.10",
            linewidth=2.0,
            edgecolor=ec,
            facecolor=fc,
            zorder=3,
        )
    )
    text = lines if isinstance(lines, str) else "\n".join(lines)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        fontweight="bold",
        color="#111827",
        linespacing=1.2,
        zorder=4,
    )
    return {
        "c": (x + w / 2, y + h / 2),
        "l": (x, y + h / 2),
        "r": (x + w, y + h / 2),
        "b": (x + w / 2, y),
        "t": (x + w / 2, y + h),
    }


def _lane(ax, x, y, w, h, title, color):
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.01,rounding_size=0.08",
            linewidth=2.0,
            edgecolor=color,
            facecolor=color,
            alpha=0.08,
            zorder=0,
        )
    )
    ax.text(
        x + 0.2,
        y + h - 0.22,
        title,
        fontsize=12,
        fontweight="bold",
        color=color,
        va="top",
        zorder=2,
    )


def _h_arrow(ax, a, b, color="#1F2937", label=None):
    ax.annotate(
        "",
        xy=(b["l"][0] - 0.04, b["l"][1]),
        xytext=(a["r"][0] + 0.04, a["r"][1]),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=2.2, mutation_scale=15),
        zorder=2,
    )
    if label:
        ax.text(
            (a["r"][0] + b["l"][0]) / 2,
            a["c"][1] + 0.2,
            label,
            ha="center",
            fontsize=9,
            color=color,
            fontweight="bold",
            zorder=5,
        )


def _v_arrow(ax, top, bottom, color="#1F2937", label=None, rad=0.0):
    ax.annotate(
        "",
        xy=(bottom["t"][0], bottom["t"][1] + 0.03),
        xytext=(top["b"][0], top["b"][1] - 0.03),
        arrowprops=dict(
            arrowstyle="-|>",
            color=color,
            lw=2.2,
            mutation_scale=15,
            connectionstyle=f"arc3,rad={rad}",
        ),
        zorder=2,
    )
    if label:
        ax.text(
            top["b"][0] + 0.2,
            (top["b"][1] + bottom["t"][1]) / 2,
            label,
            fontsize=9,
            color=color,
            fontweight="bold",
            zorder=5,
        )


# Color tokens matching Mermaid classDefs
EXIST = ("#CCFBF1", "#0F766E")  # teal — existing
PROP = ("#EDE9FE", "#7C3AED")  # purple — proposed
GATE = ("#FEF3C7", "#B45309")  # amber — quality PRs
OK = ("#DCFCE7", "#15803D")
BAD = ("#FEE2E2", "#B91C1C")
OUT = ("#FFEDD5", "#C2410C")


def render_mstf_architecture_diagram(png_path: str) -> str:
    """Full MSTF diagram with Existing vs Proposed (accurate Mermaid equivalent)."""
    fig, ax = plt.subplots(figsize=(18, 13), dpi=200)
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 13)
    ax.axis("off")
    fig.patch.set_facecolor("#FFFFFF")

    ax.text(
        9,
        12.55,
        "Forward Architecture: Multi-Surface Transaction Fabric (MSTF)",
        ha="center",
        fontsize=18,
        fontweight="bold",
        color="#1F3864",
    )
    ax.text(
        9,
        12.15,
        "Mentrix Code Red  |  Accurate Existing vs Proposed  |  Do not add ZECT Lattice",
        ha="center",
        fontsize=11,
        color="#475569",
    )

    # Legend chips
    for i, (label, fc, ec) in enumerate(
        [
            ("EXISTING", *EXIST),
            ("PROPOSED / NEW", *PROP),
            ("QUALITY PRs #118/#96", *GATE),
            ("CONTINUE", *OK),
            ("REFUSE / REPAIR", *BAD),
        ]
    ):
        x = 0.6 + i * 3.4
        ax.add_patch(
            FancyBboxPatch(
                (x, 11.55),
                3.1,
                0.4,
                boxstyle="round,pad=0.02,rounding_size=0.08",
                linewidth=1.8,
                edgecolor=ec,
                facecolor=fc,
                zorder=3,
            )
        )
        ax.text(x + 1.55, 11.75, label, ha="center", va="center", fontsize=10, fontweight="bold")

    # STEP 1 INTAKE
    _lane(ax, 0.35, 9.55, 17.3, 1.8, "STEP 1 — INTAKE", "#1D4ED8")
    jira = _box(ax, 0.7, 9.85, 2.5, 1.05, ["EXISTING", "1. Jira ticket"], *EXIST, 11)
    rfc = _box(ax, 3.7, 9.85, 2.8, 1.05, ["EXISTING", "2. RFC /", "requirements"], *EXIST, 11)
    clas = _box(
        ax,
        7.0,
        9.85,
        3.5,
        1.05,
        ["PROPOSED", "3. Scope classifier", "surfaces_required[]"],
        *PROP,
        11,
    )
    known = _box(ax, 11.1, 10.35, 6.1, 0.5, ["YES → surfaces known → continue"], *OK, 11)
    miss = _box(ax, 11.1, 9.7, 6.1, 0.5, ["NO → surface missing → REFUSE"], *BAD, 11)
    _h_arrow(ax, jira, rfc)
    _h_arrow(ax, rfc, clas)
    ax.annotate(
        "",
        xy=known["l"],
        xytext=(clas["r"][0], clas["c"][1] + 0.2),
        arrowprops=dict(arrowstyle="-|>", color=OK[1], lw=2.2, mutation_scale=14),
    )
    ax.annotate(
        "",
        xy=miss["l"],
        xytext=(clas["r"][0], clas["c"][1] - 0.2),
        arrowprops=dict(arrowstyle="-|>", color=BAD[1], lw=2.2, mutation_scale=14),
    )

    # STEP 2 + 3 side by side
    _lane(
        ax,
        0.35,
        6.85,
        8.3,
        2.5,
        "STEP 2 — KNOWLEDGE FABRIC  ·  ALL EXISTING (reuse — no Lattice)",
        EXIST[1],
    )
    _lane(
        ax,
        8.85,
        6.85,
        8.8,
        2.5,
        "STEP 3 — MSTF CONTROL PLANE  ·  ALL PROPOSED / NEW",
        PROP[1],
    )

    kb1 = _box(ax, 0.6, 7.2, 1.85, 1.45, ["EXISTING", "Repo", "Blueprints"], *EXIST, 10)
    kb2 = _box(ax, 2.6, 7.2, 1.85, 1.45, ["EXISTING", "Neo4j KG", "/ HLD"], *EXIST, 10)
    kb3 = _box(ax, 4.6, 7.2, 1.85, 1.45, ["EXISTING", "Knowledge", "Docs"], *EXIST, 10)
    kb4 = _box(ax, 6.6, 7.2, 1.8, 1.45, ["EXISTING", "Playbook", "KG"], *EXIST, 10)

    reg = _box(ax, 9.1, 8.0, 2.5, 0.95, ["PROPOSED", "4. Surface registry"], *PROP, 11)
    recipe = _box(ax, 11.95, 8.0, 2.6, 0.95, ["PROPOSED", "5. MultiSurface Recipe"], *PROP, 11)
    plan = _box(ax, 14.9, 8.0, 2.45, 0.95, ["PROPOSED", "6. Cross-repo plan"], *PROP, 11)
    refuse = _box(
        ax,
        11.95,
        7.1,
        5.4,
        0.6,
        ["PROPOSED — Refuse path (honest stop + checklist)"],
        *BAD,
        10,
    )

    ax.annotate(
        "",
        xy=reg["t"],
        xytext=(known["c"][0] - 1.5, known["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=PROP[1],
            lw=2.4,
            mutation_scale=15,
            connectionstyle="arc3,rad=0.12",
        ),
    )
    ax.text(11.3, 9.4, "continue", fontsize=9, color=PROP[1], fontweight="bold")

    ax.annotate(
        "",
        xy=(refuse["c"][0] + 1.0, refuse["t"][1]),
        xytext=(miss["c"][0], miss["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=BAD[1],
            lw=2.2,
            mutation_scale=14,
            connectionstyle="arc3,rad=-0.05",
        ),
    )

    _h_arrow(ax, reg, recipe, color=PROP[1])
    _h_arrow(ax, recipe, plan, color=PROP[1])

    # Knowledge feeds recipe — single arrow
    ax.annotate(
        "",
        xy=recipe["l"],
        xytext=(kb4["r"][0] + 0.05, kb4["c"][1]),
        arrowprops=dict(arrowstyle="-|>", color=EXIST[1], lw=3.0, mutation_scale=18),
    )
    ax.text(
        8.55,
        7.55,
        "feeds\ncontext",
        ha="center",
        fontsize=9,
        color=EXIST[1],
        fontweight="bold",
    )

    # STEP 4
    _lane(ax, 0.35, 2.85, 17.3, 3.8, "STEP 4 — EXECUTION + QUALITY GATES", OUT[1])
    ax.text(0.7, 6.2, "Surfaces:", fontsize=11, fontweight="bold", color="#7C2D12")
    for i, (label, note, fc) in enumerate(
        [
            ("NGC", "partly exists", "#86EFAC"),
            ("BPM PI", "partly exists", "#93C5FD"),
            ("CDS", "PROPOSED", "#FCD34D"),
            ("Tango", "PROPOSED", "#F9A8D4"),
        ]
    ):
        x = 2.4 + i * 2.5
        ax.add_patch(
            FancyBboxPatch(
                (x, 5.85),
                2.25,
                0.55,
                boxstyle="round,pad=0.02,rounding_size=0.08",
                linewidth=1.5,
                edgecolor="#57534E",
                facecolor=fc,
                zorder=3,
            )
        )
        ax.text(
            x + 1.125,
            6.12,
            f"{label}\n({note})",
            ha="center",
            va="center",
            fontsize=9,
            fontweight="bold",
        )

    gen = _box(
        ax,
        0.6,
        3.85,
        3.1,
        1.55,
        ["PROPOSED extend", "7. Per-surface", "generators", "NGC/BPM/CDS/Tango"],
        *PROP,
        10,
    )
    trunc = _box(ax, 4.2, 4.0, 2.9, 1.25, ["LAND / IN FLIGHT", "8. Truncation", "PR #118"], *GATE, 11)
    gates = _box(ax, 7.6, 4.0, 2.9, 1.25, ["LAND / IN FLIGHT", "9. Quality gates", "PR #96"], *GATE, 11)
    prs = _box(ax, 11.0, 4.0, 2.9, 1.25, ["EXISTING path", "10. PRs per repo", "(gated promote)"], *OUT, 11)
    score = _box(ax, 14.4, 4.0, 2.9, 1.25, ["PROPOSED", "11. Acceptance", "scorecard"], *PROP, 11)

    ax.annotate(
        "",
        xy=(gen["c"][0] + 0.6, gen["t"][1]),
        xytext=(plan["c"][0] - 0.3, plan["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=OUT[1],
            lw=2.5,
            mutation_scale=15,
            connectionstyle="arc3,rad=0.22",
        ),
    )
    ax.text(14.0, 6.55, "execute plan", fontsize=9, color=OUT[1], fontweight="bold")

    _h_arrow(ax, gen, trunc, color=OUT[1], label="then")
    _h_arrow(ax, trunc, gates, color=OUT[1], label="then")
    _h_arrow(ax, gates, prs, color=OUT[1], label="pass")
    _h_arrow(ax, prs, score, color=BAD[1], label="measure")

    ax.annotate(
        "",
        xy=(gen["c"][0], gen["b"][1] - 0.04),
        xytext=(gates["c"][0], gates["b"][1] - 0.04),
        arrowprops=dict(
            arrowstyle="-|>",
            color=BAD[1],
            lw=2.4,
            mutation_scale=15,
            connectionstyle="arc3,rad=-0.45",
        ),
    )
    ax.text(
        6.0,
        3.2,
        "FAIL → repair loop (regenerate) — never mark incomplete as done",
        ha="center",
        fontsize=11,
        color=BAD[1],
        fontweight="bold",
    )

    # STEP 5
    _lane(ax, 0.35, 0.3, 17.3, 2.3, "STEP 5 — OUTPUT / “100%” SCORECARD", BAD[1])
    o1 = _box(ax, 0.7, 0.65, 5.1, 1.4, ["Gated PRs", "only after gates pass"], *OUT, 12)
    o2 = _box(
        ax,
        6.3,
        0.65,
        5.3,
        1.4,
        ["Refuse checklist", "when CDS/Tango surface missing"],
        *BAD,
        12,
    )
    o3 = _box(
        ax,
        12.1,
        0.65,
        5.2,
        1.4,
        ["Golden-suite 100%", "complete + grounded", "+ contracted  OR  blocked"],
        *OUT,
        11,
    )
    _v_arrow(ax, score, o1, color=BAD[1])
    ax.annotate(
        "",
        xy=o2["t"],
        xytext=(refuse["c"][0], refuse["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=BAD[1],
            lw=2.0,
            mutation_scale=14,
            connectionstyle="arc3,rad=0.18",
        ),
    )
    ax.text(
        9.0,
        0.4,
        "Teal = EXISTING  |  Purple = PROPOSED  |  Amber = land PR #118/#96  |  Implementation home = Mentrix (not ZECT)",
        ha="center",
        fontsize=9,
        color="#334155",
        fontweight="bold",
    )

    fig.tight_layout(pad=0.2)
    fig.savefig(png_path, bbox_inches="tight", facecolor="white", pad_inches=0.25)
    plt.close(fig)
    return png_path


def render_mstf_happy_path_diagram(png_path: str) -> str:
    """Linear happy-path vs refuse (second Mermaid equivalent)."""
    fig, ax = plt.subplots(figsize=(16, 4.8), dpi=200)
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 4.8)
    ax.axis("off")
    fig.patch.set_facecolor("#FFFFFF")

    ax.text(
        8,
        4.45,
        "Happy path vs Refuse (same MSTF — simplified)",
        ha="center",
        fontsize=15,
        fontweight="bold",
        color="#1F3864",
    )

    a = _box(ax, 0.3, 2.55, 1.7, 1.0, ["EXISTING", "Jira"], *EXIST, 11)
    b = _box(ax, 2.4, 2.55, 2.1, 1.0, ["PROPOSED", "Classifier"], *PROP, 11)
    c = _box(ax, 5.0, 2.55, 2.6, 1.0, ["PROPOSED", "Registry +", "MultiSurface Recipe"], *PROP, 10)
    d = _box(ax, 8.05, 2.55, 2.3, 1.0, ["EXISTING", "Knowledge", "feeds context"], *EXIST, 10)
    e = _box(ax, 10.75, 2.55, 2.3, 1.0, ["PROPOSED", "Generate", "per surface"], *PROP, 10)
    f = _box(ax, 13.4, 2.55, 2.3, 1.0, ["#118 → #96", "then gated PRs"], *GATE, 10)

    _h_arrow(ax, a, b)
    _h_arrow(ax, b, c, label="surfaces known")
    _h_arrow(ax, c, d)
    _h_arrow(ax, d, e)
    _h_arrow(ax, e, f)

    refuse = _box(
        ax,
        5.0,
        0.55,
        6.0,
        1.0,
        ["NO / missing CDS·Tango → REFUSE checklist (PROPOSED)"],
        *BAD,
        11,
    )
    ax.annotate(
        "",
        xy=refuse["t"],
        xytext=(b["c"][0], b["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=BAD[1],
            lw=2.2,
            mutation_scale=14,
            connectionstyle="arc3,rad=0.25",
        ),
    )
    ax.text(3.5, 1.5, "surface missing", fontsize=9, color=BAD[1], fontweight="bold")

    ax.annotate(
        "",
        xy=(e["c"][0], e["b"][1] - 0.02),
        xytext=(f["c"][0] - 0.5, f["b"][1] - 0.02),
        arrowprops=dict(
            arrowstyle="-|>",
            color=BAD[1],
            lw=2.0,
            mutation_scale=12,
            connectionstyle="arc3,rad=-0.5",
        ),
    )
    ax.text(12.3, 1.85, "FAIL → repair", fontsize=9, color=BAD[1], fontweight="bold")

    fig.tight_layout(pad=0.2)
    fig.savefig(png_path, bbox_inches="tight", facecolor="white", pad_inches=0.2)
    plt.close(fig)
    return png_path


def main() -> tuple[str, str]:
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    os.makedirs(desktop, exist_ok=True)
    out_path = os.path.join(
        desktop, "MSTF_Multi_Surface_Transaction_Fabric_Mentrix.docx"
    )
    assets_dir = os.path.join(os.path.dirname(__file__), "..", "docs", "assets")
    os.makedirs(assets_dir, exist_ok=True)

    desktop_png = os.path.join(desktop, "MSTF_Architecture_Flow_Diagram.png")
    desktop_happy = os.path.join(desktop, "MSTF_HappyPath_vs_Refuse.png")
    diagram_png = os.path.abspath(os.path.join(assets_dir, "mstf_architecture_flow.png"))
    happy_png = os.path.abspath(os.path.join(assets_dir, "mstf_happy_path.png"))

    render_mstf_architecture_diagram(diagram_png)
    render_mstf_architecture_diagram(desktop_png)
    render_mstf_happy_path_diagram(happy_png)
    render_mstf_happy_path_diagram(desktop_happy)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    def add_heading_custom(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            set_run_font(
                run,
                size=16 if level == 1 else 13 if level == 2 else 12,
                bold=True,
                color=(31, 56, 100) if level == 1 else (47, 84, 150),
            )
        return h

    def add_para(text, bold=False, size=11, space_after=8):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        if p.runs:
            set_run_font(p.runs[0], size=11)
        else:
            set_run_font(p.add_run(text), size=11)
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            hdr.cells[i].text = h
            for p in hdr.cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        shade_header_row(hdr)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=10)
        doc.add_paragraph()
        return table

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Multi-Surface Transaction Fabric (MSTF)")
    set_run_font(r, size=22, bold=True, color=(31, 56, 100))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Forward Architecture & Implementation Plan for Mentrix")
    set_run_font(r, size=14, bold=True, color=(47, 84, 150))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Accurate Existing vs Proposed diagrams included\n"
        "Date: July 2026  |  Implementation home: Mentrix (not ZECT Lattice)"
    )
    set_run_font(r, size=10, color=(89, 89, 89))

    add_para(
        "One-liner: MSTF makes Code Red multi-system — classify Jira into NGC/BPM/CDS/Tango "
        "surfaces, generate from indexed Blueprint + Knowledge + Playbooks, fail closed with "
        "quality gates (PRs #118/#96), measure “100%” on a golden suite — all inside Mentrix.",
        bold=True,
    )

    add_heading_custom("1. Problem Statement", 1)
    add_para(
        "Mentrix already succeeds when a Jira maps to an exemplar-similar transaction and "
        "changes stay on the NGC / BPM PI recipe path (e.g. Authorized Signatory). It fails or "
        "is bypassed when the same ticket also needs CDS and/or Tango work. Historical POC "
        "accuracy was ~50%. PRs #118/#96 harden quality; they do not alone add CDS/Tango coverage."
    )
    add_table(
        ["Situation", "Reality"],
        [
            ["Authorized Signatory + NGC-like", "Works today (EXISTING recipe path)"],
            ["Same program + CDS / Tango", "Gap — needs PROPOSED MSTF surfaces"],
            ["PRs #118 / #96", "Quality kernel to LAND (not domain coverage)"],
        ],
    )

    add_heading_custom("2. Design Principles", 1)
    for t in [
        "No Lattice in Mentrix — reuse Blueprint, Neo4j KG, KnowledgeDocs, Playbook KG, TransactionRecipe.",
        "Refuse > hallucinate — missing surface → hard-stop with checklist.",
        "AI-agnostic control plane — models pluggable; truth in KB/playbook/recipe/gates.",
        "Capability-scoped repo selection — not unconstrained “LLM picks any repo.”",
        "“100%” = closed loop — complete, grounded, contracted, or blocked.",
        "Extend, don’t rewrite — wrap orchestrator → planner → codegen → PR.",
    ]:
        add_bullet(t)

    add_heading_custom("3. Existing vs Proposed (accurate map)", 1)
    add_para(
        "This table matches the architecture diagrams below and the Mermaid Existing/Proposed view.",
        bold=True,
    )
    add_table(
        ["Layer / component", "Status", "Notes"],
        [
            ["Jira ticket → RFC / requirements", "EXISTING", "Orchestrator intake already exists"],
            ["Scope classifier (surfaces_required)", "PROPOSED", "Formal NGC/BPM/CDS/Tango + refuse"],
            ["Repo Blueprints", "EXISTING", "blueprint-generator"],
            ["Neo4j KG / HLD", "EXISTING", "orchestrator indexing"],
            ["Knowledge Docs", "EXISTING", "repo walk + KnowledgeDoc nodes"],
            ["Playbook KG bundles (agent-index.json)", "EXISTING", "includes Playbook node label"],
            ["Surface registry", "PROPOSED", "MSTF control plane"],
            ["MultiSurface Recipe + cross-repo file plan", "PROPOSED", "extends TransactionRecipe"],
            ["Refuse path (honest stop)", "PROPOSED", "no fake codegen when surface missing"],
            ["NGC / BPM PI generators (template-like)", "EXISTING (partial)", "Authorized Signatory proven"],
            ["CDS / Tango per-surface generators", "PROPOSED", "needs index + recipe + skills"],
            ["Truncation PR #118 (Mentrix-common)", "LAND / IN FLIGHT", "generate_with_status"],
            ["Quality gates PR #96 (code-generator)", "LAND / IN FLIGHT", "AC, grounding, missing-LLD, etc."],
            ["PRs per repo path", "EXISTING", "promote/create PR machinery"],
            ["Acceptance scorecard / golden suite 100%", "PROPOSED", "closed-loop metric"],
        ],
    )

    add_heading_custom("4. Target Architecture — diagrams", 1)
    add_heading_custom("4.1 Full MSTF flow (Existing vs Proposed)", 2)
    add_para(
        "Figure 1 is the accurate Mermaid-equivalent architecture: teal = EXISTING, purple = PROPOSED, "
        "amber = land quality PRs, green = continue, red = refuse/repair.",
        bold=True,
    )
    doc.add_picture(desktop_png, width=Inches(7.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(
        "Figure 1. MSTF Forward Architecture — Existing vs Proposed (steps 1→11)"
    )
    set_run_font(r, size=9, bold=True, color=(71, 85, 105))

    add_heading_custom("4.2 Happy path vs Refuse (simplified)", 2)
    add_para(
        "Figure 2 is the linear view: known surfaces continue through knowledge + generate + gates; "
        "missing CDS/Tango surfaces take the refuse path."
    )
    doc.add_picture(desktop_happy, width=Inches(7.2))
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap2.add_run("Figure 2. Happy path vs Refuse checklist")
    set_run_font(r, size=9, bold=True, color=(71, 85, 105))

    add_para(
        "Standalone PNGs on Desktop for projector demos: "
        "MSTF_Architecture_Flow_Diagram.png and MSTF_HappyPath_vs_Refuse.png"
    )

    add_para("How to walk Figure 1 in a demo:", bold=True)
    add_bullet("Step 1 EXISTING intake (Jira→RFC) then PROPOSED classifier → YES continue / NO refuse.")
    add_bullet("Step 2 ALL EXISTING knowledge fabric feeds context (no Lattice).")
    add_bullet("Step 3 ALL PROPOSED control plane: registry → MultiSurface Recipe → cross-repo plan.")
    add_bullet("Step 4 Generate per surface → LAND #118 → LAND #96 → PRs → PROPOSED scorecard; FAIL repairs.")
    add_bullet("Step 5 Outputs: gated PRs, refuse checklist, or golden-suite 100% definition.")

    add_heading_custom("5. Quality spine (PRs to land)", 1)
    add_table(
        ["PR", "Repo", "Role"],
        [
            ["#118", "Mentrix-common", "Truncation detection (generate_with_status / finish_reason)"],
            [
                "#96",
                "Mentrix-code-generator",
                "Continuation, AC verifier, coverage, invented-API xref, missing-LLD, manifest refs",
            ],
        ],
    )
    add_bullet("https://github.com/zinnia/Mentrix-common/pull/118")
    add_bullet("https://github.com/zinnia/Mentrix-code-generator/pull/96")

    add_heading_custom("6. Implementation phases", 1)
    add_table(
        ["Phase", "Outcome", "Exit criteria"],
        [
            ["P0", "Land #118 + #96; NGC goldens", "Suite #1–2 pass"],
            ["P1", "Classifier + refuse", "Missing CDS/Tango refuses honestly"],
            ["P2", "Index CDS+Tango; ship KB/Playbooks", "Retrieval hits those docs"],
            ["P3", "MultiSurfaceRecipe v1 pilot ticket", "One CDS+Tango E2E assisted"],
            ["P4", "Extend #96 gates to new surfaces", "Blocking gates on CDS/Tango"],
            ["P5", "Adoption metrics", "% Code Red via bot measured"],
        ],
    )
    add_para("P0 alone ≠ Anubhav gap closed. P1–P4 close CDS/Tango.", bold=True)

    add_heading_custom("7. Component map (where to build)", 1)
    add_table(
        ["Concern", "Home", "Status"],
        [
            ["Jira + RFC", "Mentrix-orchestrator", "EXISTING"],
            ["Classifier + refuse", "orchestrator + common", "PROPOSED"],
            ["Surface registry", "common / Mongo", "PROPOSED"],
            ["Blueprint / KG / Knowledge ingest", "blueprint-generator + indexer", "EXISTING"],
            ["Plan", "Mentrix-planner", "EXISTING + extend"],
            ["MultiSurfaceRecipe + codegen", "code-generator + common", "PROPOSED extend"],
            ["Truncation #118", "Mentrix-common", "LAND"],
            ["Quality gates #96", "code-generator", "LAND"],
        ],
    )

    add_heading_custom("8. ZECT boundary", 1)
    add_table(
        ["Question", "Decision"],
        [
            ["Build MSTF in Mentrix?", "YES"],
            ["Build MSTF in ZECT?", "NO for Code Red gap"],
            ["Add Lattice to Mentrix?", "NO — use Blueprint + Neo4j + KB + Playbooks"],
            ["Later Mentrix calls Mentrix?", "Optional integration only"],
        ],
    )

    add_heading_custom("9. Immediate asks", 1)
    add_bullet("Confirm merge path for #118 and #96.")
    add_bullet("Lasya / Anubhav: one-page repo list per surface (NGC, CDS, Tango).")
    add_bullet("Siddartha: refresh + index those repos.")
    add_bullet("Pick one paused CDS+Tango Code Red ticket as P3 pilot.")
    add_bullet("Agree scorecard wording: leverage Mentrix now on NGC-like; CDS/Tango via MSTF phases.")

    add_heading_custom("10. Demo talk track (5–7 min)", 1)
    add_table(
        ["Step", "Say this"],
        [
            ["1", "Works for NGC/BPM exemplar-like; CDS/Tango still manual; old POC ~50%."],
            ["2", "Teal boxes already exist; purple boxes are what we propose to build."],
            ["3", "Amber boxes = land PRs #118/#96 (real code quality, not docs)."],
            ["4", "Will not add ZECT Lattice into Mentrix."],
            ["5", "Refuse when surface missing — never fake CDS/Tango success."],
            ["6", "Ask: merge PRs → glossary → index → one pilot → golden suite."],
        ],
    )

    foot = doc.add_paragraph()
    r = foot.add_run(
        "Document owner: Mentrix / platform engineering. Diagrams match Existing vs Proposed Mermaid. "
        "Update surface registry as CDS/Tango repos are confirmed."
    )
    set_run_font(r, size=9, color=(89, 89, 89))

    doc.save(out_path)
    return out_path, desktop_png


if __name__ == "__main__":
    path, png = main()
    print(path)
    print(png)
    happy = os.path.join(os.path.expanduser("~"), "Desktop", "MSTF_HappyPath_vs_Refuse.png")
    print(happy)
    print("docx_size_bytes", os.path.getsize(path))
    print("png_size_bytes", os.path.getsize(png))
    print("happy_size_bytes", os.path.getsize(happy))
